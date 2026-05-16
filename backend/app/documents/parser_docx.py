from __future__ import annotations

from pathlib import Path

from backend.app.documents.parser_pdf import ParsedDocument, ParsedPage


def parse_docx(path: Path) -> ParsedDocument:
    import docx

    document = docx.Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return ParsedDocument(pages=[ParsedPage(page_number=1, text=text)])

